#!/usr/bin/env python3.12
"""デジラボビューティー定款をエステティック協会書式で作成"""
import sys
sys.path.insert(0, '/usr/lib/python3/dist-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ設定
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# スタイルヘルパー
def set_font(run, size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'MS明朝'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS明朝')

def add_title(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size, bold)
    return p

def add_chapter(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 10.5, False)
    return p

def add_article_title(text):
    """（○○）スタイルの小見出し"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, 10.5, False)
    return p

def add_article(num, text):
    """第X条 本文"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'第 {num} 条　{text}')
    set_font(run, 10.5, False)
    return p

def add_para(num, text):
    """2　本文（項番号付き）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(f'{num}　{text}')
    set_font(run, 10.5, False)
    return p

def add_item(num, text):
    """（1）　箇条書き"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f'（{num}）　{text}')
    set_font(run, 10.5, False)
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    set_font(run, 6, False)
    return p

def add_appendix_heading(num, text):
    """附則番号付き見出し"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{num}　{text}')
    set_font(run, 10.5, False)
    return p

def add_normal(text, indent=0):
    """普通のテキスト"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, 10.5, False)
    return p

# ============================================================
# タイトル
# ============================================================
add_title('一般社団法人デジラボビューティー　定款', size=14, bold=True)
add_blank()

# ============================================================
# 第 1 章　総則
# ============================================================
add_chapter('第 1 章　総則')

add_article_title('（名称）')
add_article('1', 'この法人は、一般社団法人デジラボビューティーと称し、英文では DigiLab Beauty と表示する。')

add_article_title('（事務所）')
add_article('2', 'この法人は、主たる事務所を【○○】に置く。')
add_para('2', 'この法人は、理事会の決議によって、従たる事務所を必要な地に置くことができる。')

# ============================================================
# 第 2 章　目的及び事業
# ============================================================
add_chapter('第 2 章　目的及び事業')

add_article_title('（目的）')
add_article('3', 'この法人は、人工知能（AI）及びデジタル技術の健全な活用を通じて、美容及びエステティック分野における品質、安全性並びに倫理に関する標準を確立し、もって事業者及び従事者の専門性の向上並びに消費者の信頼の確保に寄与し、我が国及びアジアにおける美容産業の持続的な発展に貢献することを目的とする。')

add_article_title('（事業）')
add_article('4', 'この法人は、前条の目的を達成するため、次の事業を行う。')
add_item('1', '美容及びエステティック分野における AI・デジタル技術活用に関する品質標準の調査研究、策定及び普及')
add_item('2', '美容関連の事業者及び従事者に対する認証制度並びに検定制度の企画、運営及び認証の付与')
add_item('3', '前号に関する教育、研修並びに教科書及び教材の編集及び発行')
add_item('4', '美容産業に関する調査研究、白書及びガイドラインの策定並びに政策提言')
add_item('5', '会員相互の交流、情報の提供及び広報並びに普及啓発')
add_item('6', '国内外の関係機関、団体及び行政との連携並びに国際標準化への参画')
add_item('7', '前各号に附帯し、又は関連する一切の事業')
add_para('2', '前項の事業は、本邦及び海外において行うものとする。')

# ============================================================
# 第 3 章　社員及び会員
# ============================================================
add_chapter('第 3 章　社員及び会員')

add_article_title('（法人の構成員）')
add_article('5', 'この法人の社員は、この法人の目的に賛同して入社した個人又は団体であって、次条の規定により入社した者とする。')
add_para('2', 'この法人の事業を賛助するため、社員以外に賛助会員を置くことができる。賛助会員は、社員総会における議決権を有しない。賛助会員に関する事項は、社員総会の決議により別に定める会員規程による。')

add_article_title('（入社）')
add_article('6', '社員になろうとする者は、理事会の定める入社申込書を提出し、理事会の承認を受けなければならない。')

add_article_title('（経費等の負担）')
add_article('7', '社員及び賛助会員は、この法人の事業活動に経常的に生じる費用に充てるため、社員総会において別に定める額の経費を負担するものとする。')

add_article_title('（任意退社）')
add_article('8', '社員は、いつでも退社することができる。ただし、退社しようとするときは、あらかじめ書面又は電磁的方法をもってこの法人に通知しなければならない。')

add_article_title('（除名）')
add_article('9', '社員が次のいずれかに該当するに至ったときは、社員総会の決議によって当該社員を除名することができる。この場合、その社員に対し、社員総会の日から 1 週間前までにその旨を通知し、かつ、社員総会において弁明の機会を与えなければならない。')
add_item('1', 'この定款その他の規則に違反したとき。')
add_item('2', 'この法人の名誉を傷つけ、又は目的に反する行為をしたとき。')
add_item('3', 'その他除名すべき正当な事由があるとき。')

add_article_title('（社員の資格の喪失）')
add_article('10', '前 2 条のほか、社員は、次のいずれかに該当するに至ったときは、その資格を喪失する。')
add_item('1', '第 7 条の支払義務を 2 年以上履行しなかったとき。')
add_item('2', '総社員が同意したとき。')
add_item('3', '当該社員が死亡し、若しくは解散し、又は後見開始若しくは保佐開始の審判を受けたとき。')

# ============================================================
# 第 4 章　社員総会
# ============================================================
add_chapter('第 4 章　社員総会')

add_article_title('（構成）')
add_article('11', '社員総会は、すべての社員をもって構成する。')

add_article_title('（権限）')
add_article('12', '社員総会は、次の事項について決議する。')
add_item('1', '社員の除名')
add_item('2', '理事及び監事の選任又は解任')
add_item('3', '理事及び監事の報酬等の額')
add_item('4', '貸借対照表及び損益計算書（正味財産増減計算書）の承認')
add_item('5', '定款の変更')
add_item('6', '解散及び残余財産の処分')
add_item('7', 'その他社員総会で決議するものとして法令又はこの定款で定められた事項')

add_article_title('（開催）')
add_article('13', '社員総会は、定時社員総会として毎事業年度終了後 3 か月以内に 1 回開催するほか、必要がある場合に臨時社員総会を開催する。')

add_article_title('（招集）')
add_article('14', '社員総会は、法令に別段の定めがある場合を除き、理事会の決議に基づき代表理事が招集する。')
add_para('2', '代表理事に事故等があるときは、あらかじめ理事会が定めた順序により、他の理事が社員総会を招集する。')

add_article_title('（議長）')
add_article('15', '社員総会の議長は、代表理事がこれに当たる。代表理事に事故等があるときは、当該社員総会において出席した社員の中から選出する。')

add_article_title('（議決権）')
add_article('16', '社員総会における議決権は、社員 1 名につき 1 個とする。')

add_article_title('（決議）')
add_article('17', '社員総会の決議は、法令又はこの定款に別段の定めがある場合を除き、総社員の議決権の過半数を有する社員が出席し、出席した当該社員の議決権の過半数をもって行う。')
add_para('2', '前項の規定にかかわらず、次の決議は、総社員の半数以上であって、総社員の議決権の 3 分の 2 以上に当たる多数をもって行う。')
add_item('1', '社員の除名')
add_item('2', '監事の解任')
add_item('3', '定款の変更')
add_item('4', '解散')
add_item('5', 'その他法令で定められた事項')

add_article_title('（議事録）')
add_article('18', '社員総会の議事については、法令で定めるところにより議事録を作成し、議長及び出席した理事は、これに署名し、又は記名押印する。')

# ============================================================
# 第 5 章　役員
# ============================================================
add_chapter('第 5 章　役員')

add_article_title('（役員の設置）')
add_article('19', 'この法人に、次の役員を置く。')
add_item('1', '理事　3 名以上 10 名以内')
add_item('2', '監事　1 名以上 2 名以内')
add_para('2', '理事のうち 1 名を代表理事とする。')
add_para('3', '代表理事以外の理事のうち、若干名を業務執行理事とすることができる。')

add_article_title('（役員の選任）')
add_article('20', '理事及び監事は、社員総会の決議によって選任する。')
add_para('2', '代表理事及び業務執行理事は、理事会の決議によって理事の中から選定する。')
add_para('3', '各理事について、当該理事及びその配偶者又は三親等以内の親族その他当該理事と特別の関係がある者である理事の合計数は、理事の総数の 3 分の 1 を超えてはならない。監事についても同様とする。')

add_article_title('（理事の職務及び権限）')
add_article('21', '理事は、理事会を構成し、法令及びこの定款で定めるところにより、職務を執行する。')
add_para('2', '代表理事は、法令及びこの定款で定めるところにより、この法人を代表し、その業務を執行する。')
add_para('3', '代表理事及び業務執行理事は、毎事業年度に 4 か月を超える間隔で 2 回以上、自己の職務の執行の状況を理事会に報告しなければならない。')

add_article_title('（監事の職務及び権限）')
add_article('22', '監事は、理事の職務の執行を監査し、法令で定めるところにより監査報告を作成する。')
add_para('2', '監事は、いつでも、理事及び使用人に対して事業の報告を求め、この法人の業務及び財産の状況の調査をすることができる。')

add_article_title('（役員の任期）')
add_article('23', '理事の任期は、選任後 2 年以内に終了する事業年度のうち最終のものに関する定時社員総会の終結の時までとする。ただし、再任を妨げない。')
add_para('2', '監事の任期は、選任後 4 年以内に終了する事業年度のうち最終のものに関する定時社員総会の終結の時までとする。ただし、再任を妨げない。')
add_para('3', '補欠として選任された理事又は監事の任期は、前任者の任期の満了する時までとする。')
add_para('4', '理事又は監事は、第 19 条に定める定数に足りなくなるときは、任期の満了又は辞任により退任した後も、新たに選任された者が就任するまで、なお理事又は監事としての権利義務を有する。')

add_article_title('（役員の解任）')
add_article('24', '理事及び監事は、社員総会の決議によって解任することができる。ただし、監事を解任する決議は、第 17 条第 2 項の決議によらなければならない。')

add_article_title('（役員の報酬等）')
add_article('25', '理事及び監事の報酬、賞与その他の職務執行の対価としてこの法人から受ける財産上の利益は、社員総会の決議によって定める。')

# ============================================================
# 第 6 章　理事会
# ============================================================
add_chapter('第 6 章　理事会')

add_article_title('（構成）')
add_article('26', 'この法人に理事会を置く。')
add_para('2', '理事会は、すべての理事をもって構成する。')

add_article_title('（権限）')
add_article('27', '理事会は、次の職務を行う。')
add_item('1', 'この法人の業務執行の決定')
add_item('2', '理事の職務の執行の監督')
add_item('3', '代表理事及び業務執行理事の選定及び解職')

add_article_title('（招集）')
add_article('28', '理事会は、代表理事が招集する。代表理事に事故等があるときは、あらかじめ理事会が定めた順序により、他の理事が招集する。')

add_article_title('（決議）')
add_article('29', '理事会の決議は、決議について特別の利害関係を有する理事を除く理事の過半数が出席し、その過半数をもって行う。')
add_para('2', '前項の規定にかかわらず、法人法第 96 条の要件を満たすときは、理事会の決議があったものとみなすことができる。')

add_article_title('（議事録）')
add_article('30', '理事会の議事については、法令で定めるところにより議事録を作成し、出席した代表理事及び監事は、これに署名し、又は記名押印する。')

# ============================================================
# 第 7 章　資産及び会計
# ============================================================
add_chapter('第 7 章　資産及び会計')

add_article_title('（事業年度）')
add_article('31', 'この法人の事業年度は、毎年【月　日】から翌年【月　日】までの年 1 期とする。')

add_article_title('（事業計画及び収支予算）')
add_article('32', 'この法人の事業計画書及び収支予算書は、毎事業年度の開始の日の前日までに代表理事が作成し、理事会の承認を受けなければならない。これを変更する場合も同様とする。')

add_article_title('（事業報告及び決算）')
add_article('33', 'この法人の事業報告及び決算については、毎事業年度終了後、代表理事が次の書類を作成し、監事の監査を受けた上で、理事会の承認を経て、定時社員総会に提出し、又は提供しなければならない。')
add_item('1', '事業報告及びその附属明細書')
add_item('2', '貸借対照表')
add_item('3', '損益計算書（正味財産増減計算書）')
add_para('2', '前項の規定により承認された書類のうち、第 1 号の事業報告については定時社員総会に報告し、第 2 号及び第 3 号については定時社員総会の承認を受けなければならない。')

add_article_title('（剰余金）')
add_article('34', 'この法人は、剰余金の分配を行うことができない。')

# ============================================================
# 第 8 章　定款の変更及び解散
# ============================================================
add_chapter('第 8 章　定款の変更及び解散')

add_article_title('（定款の変更）')
add_article('35', 'この定款は、社員総会において、第 17 条第 2 項の決議によって変更することができる。')

add_article_title('（解散）')
add_article('36', 'この法人は、社員総会の決議その他法令で定められた事由により解散する。')

add_article_title('（残余財産の帰属）')
add_article('37', 'この法人が清算をする場合において有する残余財産は、社員総会の決議を経て、公益社団法人及び公益財団法人の認定等に関する法律第 5 条第 17 号に掲げる法人又は国若しくは地方公共団体に贈与するものとする。')

# ============================================================
# 第 9 章　公告の方法
# ============================================================
add_chapter('第 9 章　公告の方法')

add_article_title('（公告の方法）')
add_article('38', 'この法人の公告は、電子公告により行う。ただし、事故その他やむを得ない事由によって電子公告による公告をすることができない場合は、官報に掲載する方法による。')

# ============================================================
# 附則
# ============================================================
add_blank()
add_chapter('附　則')

add_appendix_heading('1', '設立時社員の氏名及び住所')
add_normal('この法人の設立時社員の氏名及び住所は、次のとおりである。', indent=0.5)
add_item('1', '設立時社員　氏名：【鎌田　麻央】　住所：【愛媛県松山市〇〇】')
add_item('2', '設立時社員　氏名：【　　　　　　】　住所：【　　　　　　　　　　　　】')

add_appendix_heading('2', '設立時の役員')
add_normal('この法人の設立時の役員は、次のとおりである。', indent=0.5)
add_item('1', '設立時理事　【鎌田　麻央】　【　　　　　　】　【　　　　　　】')
add_item('2', '設立時代表理事　【鎌田　麻央】')
add_item('3', '設立時監事　【　　　　　　】')

add_appendix_heading('3', '最初の事業年度')
add_normal('この法人の最初の事業年度は、この法人の成立の日から【令和〇年〇月〇日】までとする。', indent=0.5)

add_appendix_heading('4', '設立時の主たる事務所')
add_normal('この法人の設立時の主たる事務所は、【〇〇】に置く。', indent=0.5)

add_appendix_heading('5', '法令の準拠')
add_normal('この定款に定めのない事項は、すべて一般社団法人及び一般財団法人に関する法律その他の法令の定めるところによる。', indent=0.5)

add_blank()
add_normal('以上、一般社団法人デジラボビューティーを設立するため、設立時社員が定款を作成し、次に記名押印する。')
add_blank()

p = doc.add_paragraph()
run = p.add_run('令和　　年　　月　　日')
set_font(run, 10.5, False)

add_blank()

p = doc.add_paragraph()
run = p.add_run('設立時社員　　住所：【　　　　　　　　　　　　　　　　】　氏名：【　　　　　　　　　　】　　　　　　　　㊞')
set_font(run, 10.5, False)

p = doc.add_paragraph()
run = p.add_run('設立時社員　　住所：【　　　　　　　　　　　　　　　　】　氏名：【　　　　　　　　　　】　　　　　　　　㊞')
set_font(run, 10.5, False)

out_path = '/home/user/xenomao/デジラボビューティー定款_v2.0.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
